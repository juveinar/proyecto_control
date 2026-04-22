from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.core.paginator import Paginator
from django.db.models import Q, Count
from django.utils import timezone
from datetime import datetime, date
import json
import time
import os
import io
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Proyecto, Evento, ProyectoFase, Contacto, ControlProyectosInventario

@require_http_methods(["GET", "POST"])
def login_view(request):
    """Vista de login"""
    if request.user.is_authenticated:
        return redirect('index')

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        if not username or not password:
            messages.error(request, 'Por favor, ingresa usuario y contraseña.')
        else:
            user = authenticate(request, username=username, password=password)
            if user is not None:
                if user.is_active:
                    login(request, user)
                    return redirect('index')
                else:
                    messages.error(request, 'Tu cuenta está desactivada.')
            else:
                messages.error(request, 'Usuario o contraseña incorrectos.')

    return render(request, 'login.html')


@login_required
def logout_view(request):
    """Vista de logout"""
    logout(request)
    return redirect('login')


@login_required
def index(request):
    """Vista principal"""
    return render(request, 'index.html', {'current_user': request.user})


# ==================== API DE PROYECTOS ====================

@require_http_methods(["GET", "POST"])
@login_required
def api_projects(request):
    """Obtener todos los proyectos o crear uno nuevo"""
    if request.method == 'GET':
        proyectos = Proyecto.objects.select_related('contacto').all().order_by('-start', 'project')

        # Convertir a lista de diccionarios
        proyectos_list = []
        for p in proyectos:
            # Obtener la fase más reciente
            fase_actual = p.fases.order_by('-fecha', '-created_at').first()
            if fase_actual:
                fase_str = f"{fase_actual.get_fase_display()} ({fase_actual.fecha.strftime('%Y-%m-%d')})"
            elif p.estado == 'En Curso':
                fase_str = "Despliegue (No registrado)"
            else:
                fase_str = "Sin Fase"

            # Obtener información del contacto
            contacto_info = ''
            if p.contacto:
                contacto_info = p.contacto.nombre
                if p.contacto.correo:
                    contacto_info += f" ({p.contacto.correo})"

            proyecto_dict = {
                'Id Project': p.id_project,
                'RF': p.rf or '',
                'Project': p.project or '',
                'Project Leader': p.project_leader or '',
                'Estado': p.estado or '',
                'Fase': fase_str,  # Nuevo campo para la tabla
                'Start': p.start.strftime('%Y-%m-%d') if p.start else None,
                'Finish': p.finish.strftime('%Y-%m-%d') if p.finish else None,
                'Computo': p.computo or '',
                'NTP': p.ntp or '',
                'SCAN': p.scan or '',
                'RESUELVE POR NOMBRE': p.resuelve_por_nombre or '',
                'Antivirus': p.antivirus or '',
                'CONFIG BACKUP': p.config_backup or '',
                'MONITOREO NAGIOS': p.monitoreo_nagios or '',
                'MONITOREO ELASTIC': p.monitoreo_elastic or '',
                'UCMDB': p.ucmdb or '',
                'CONECTIVIDAD AWX 172.18.90.250 (SOLO UNIX)': p.conectividad_awx or '',
                'CAMBIO PASO OPERACIÓN (OLA)': p.cambio_paso_operacion_ola or '',
                'Base de Datos': p.base_de_datos or '',
                'Balanceo': p.balanceo or '',
                'Backup': p.backup or '',
                'CHECK AV': p.check_av or '',
                'CANTIDAD MAQUINAS': p.cantidad_maquinas or '',
                'COD SERV_HOSTNAME': p.cod_serv_hostname or '',
                'PLATAFORMA': p.plataforma or '',
                'SO': p.so or '',
                'WINDOWS LICENCIA ACTIVADA': p.windows_licencia_activada or '',
                'DOMINIO': p.dominio or '',
                'PLATAFORMA BACKUP': p.plataforma_backup or '',
                'PROVEEDOR': p.proveedor or '',
                'COMUNIDAD SNMP': p.comunidad_snmp or '',
                'FGN 172.22.16.93': p.fgn_172_22_16_93 or '',
                'RT': p.rt or '',
                'SERVICIO': p.servicio or '',
                'OBSERVACIONES': p.observaciones or '',
                'CONTACTO': contacto_info,
                'CONTACTO_ID': p.contacto.id if p.contacto else None,
            }
            proyectos_list.append(proyecto_dict)

        return JsonResponse(proyectos_list, safe=False)

    elif request.method == 'POST':
        # Llamar a la función api_projects_add para crear un proyecto
        return api_projects_add(request)

    else:
        return JsonResponse({'error': 'Método no permitido'}, status=405)


@require_http_methods(["GET"])
@login_required
def api_projects_stats(request):
    """Obtener estadísticas de proyectos por mes"""
    year_str = request.GET.get('year')
    year = None
    if year_str:
        try:
            year = int(year_str)
        except (ValueError, TypeError):
            year = None

    proyectos = Proyecto.objects.exclude(start__isnull=True)

    if year:
        proyectos = proyectos.filter(start__year=year)

    # Agrupar por mes
    monthly_counts = [0] * 12
    for proyecto in proyectos:
        if proyecto.start:
            month = proyecto.start.month
            monthly_counts[month - 1] += 1

    labels = ['E', 'F', 'M', 'A', 'M', 'J', 'J', 'A', 'S', 'O', 'N', 'D']
    full_labels = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
                   'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']

    return JsonResponse({
        'labels': labels,
        'data': monthly_counts,
        'full_labels': full_labels
    })


@require_http_methods(["POST"])
@login_required
def api_projects_add(request):
    """Agregar un nuevo proyecto"""
    try:
        data = json.loads(request.body)

        # Campos que deben inicializarse con 'Pendiente'
        pendiente_fields = {
            'NTP': 'ntp',
            'SCAN': 'scan',
            'RESUELVE POR NOMBRE': 'resuelve_por_nombre',
            'Antivirus': 'antivirus',  # clave alineada con el frontend
            'CONFIG BACKUP': 'config_backup',
            'MONITOREO NAGIOS': 'monitoreo_nagios',
            'MONITOREO ELASTIC': 'monitoreo_elastic',
            'UCMDB': 'ucmdb',
            'CONECTIVIDAD AWX 172.18.90.250 (SOLO UNIX)': 'conectividad_awx',
            'CAMBIO PASO OPERACIÓN (OLA)': 'cambio_paso_operacion_ola',
        }

        # Crear el proyecto
        proyecto = Proyecto()

        # Mapear campos principales
        if 'Id Project' in data:
            proyecto.id_project = data['Id Project']
        if 'RF' in data:
            proyecto.rf = data['RF']
        if 'Project' in data:
            proyecto.project = data['Project']
        if 'Project Leader' in data:
            proyecto.project_leader = data['Project Leader']
        if 'Estado' in data:
            proyecto.estado = data['Estado']

        # Fechas
        for field_name, model_field in [
            ('Start', 'start'),
            ('Finish', 'finish'),
        ]:
            if field_name in data and data[field_name]:
                try:
                    date_value = datetime.strptime(data[field_name], '%Y-%m-%d').date()
                    setattr(proyecto, model_field, date_value)
                except (ValueError, TypeError):
                    pass

        # Cómputo
        if 'Computo' in data:
            proyecto.computo = data['Computo']

        # Campos pendientes
        for excel_field, model_field in pendiente_fields.items():
            value = data.get(excel_field, 'Pendiente')
            if not value or (isinstance(value, str) and value.strip() == ''):
                value = 'Pendiente'
            setattr(proyecto, model_field, value)

        # Otros campos
        if 'Base de Datos' in data:
            proyecto.base_de_datos = data['Base de Datos']
        if 'Balanceo' in data:
            proyecto.balanceo = data['Balanceo']
        if 'Backup' in data:
            proyecto.backup = data['Backup']
        if 'CHECK AV' in data:
            proyecto.check_av = data['CHECK AV']

        # Campos adicionales
        campos_adicionales = {
            'CONTACTO': 'contacto_id',  # Campo para contacto
            'CANTIDAD MAQUINAS': 'cantidad_maquinas',
            'COD SERV_HOSTNAME': 'cod_serv_hostname',
            'PLATAFORMA': 'plataforma',
            'SO': 'so',
            'WINDOWS LICENCIA ACTIVADA': 'windows_licencia_activada',
            'DOMINIO': 'dominio',
            'PLATAFORMA BACKUP': 'plataforma_backup',
            'PROVEEDOR': 'proveedor',
            'COMUNIDAD SNMP': 'comunidad_snmp',
            'FGN 172.22.16.93': 'fgn_172_22_16_93',
            'RT': 'rt',
            'SERVICIO': 'servicio',
            'OBSERVACIONES': 'observaciones',
        }

        for excel_field, model_field in campos_adicionales.items():
            if excel_field in data:
                setattr(proyecto, model_field, data[excel_field] or None)

        proyecto.save()

        # Refrescar relación contacto para asegurar datos actualizados
        if proyecto.contacto_id:
            proyecto.refresh_from_db(fields=['contacto'])

        # 1. Requerimiento: Crear automáticamente fase Despliegue al crear proyecto
        try:
            ProyectoFase.objects.create(
                proyecto=proyecto,
                fase='DESPLIEGUE',
                fecha=timezone.now().date()
            )
        except Exception as e:
            print(f"Error creando fase inicial: {e}")

        # Retornar el proyecto creado en formato JSON
        contacto_info = ''
        if proyecto.contacto:
            contacto_info = proyecto.contacto.nombre
            if proyecto.contacto.correo:
                contacto_info += f" ({proyecto.contacto.correo})"

        proyecto_dict = {
            'Id Project': proyecto.id_project,
            'RF': proyecto.rf or '',
            'Project': proyecto.project or '',
            'Project Leader': proyecto.project_leader or '',
            'Estado': proyecto.estado or '',
            'Fase': f"Despliegue ({timezone.now().strftime('%Y-%m-%d')})",
            'Start': proyecto.start.strftime('%Y-%m-%d') if proyecto.start else None,
            'Finish': proyecto.finish.strftime('%Y-%m-%d') if proyecto.finish else None,
            'Computo': proyecto.computo or '',
            'CANTIDAD MAQUINAS': proyecto.cantidad_maquinas or '',
            'COD SERV_HOSTNAME': proyecto.cod_serv_hostname or '',
            'PLATAFORMA': proyecto.plataforma or '',
            'SO': proyecto.so or '',
            'WINDOWS LICENCIA ACTIVADA': proyecto.windows_licencia_activada or '',
            'DOMINIO': proyecto.dominio or '',
            'PLATAFORMA BACKUP': proyecto.plataforma_backup or '',
            'PROVEEDOR': proyecto.proveedor or '',
            'COMUNIDAD SNMP': proyecto.comunidad_snmp or '',
            'FGN 172.22.16.93': proyecto.fgn_172_22_16_93 or '',
            'RT': proyecto.rt or '',
            'SERVICIO': proyecto.servicio or '',
            'OBSERVACIONES': proyecto.observaciones or '',
            'CONTACTO': contacto_info,
            'CONTACTO_ID': proyecto.contacto.id if proyecto.contacto else None,
        }

        return JsonResponse(proyecto_dict, status=201)

    except Exception as e:
        return JsonResponse({'error': f'No se pudo guardar el nuevo proyecto: {str(e)}'}, status=500)


@require_http_methods(["PUT"])
@login_required
def api_projects_update(request, project_id):
    """Actualizar un proyecto existente"""
    if request.method != 'PUT':
        return JsonResponse({'error': 'Método no permitido'}, status=405)

    try:
        proyecto = Proyecto.objects.select_related('contacto').get(id_project=project_id)
        data = json.loads(request.body)

        # Actualizar campos principales
        if 'RF' in data:
            proyecto.rf = data['RF']
        if 'Project' in data:
            proyecto.project = data['Project']
        if 'Project Leader' in data:
            proyecto.project_leader = data['Project Leader']
        if 'Estado' in data:
            proyecto.estado = data['Estado']

        # Fechas
        for field_name, model_field in [
            ('Start', 'start'),
            ('Finish', 'finish'),
        ]:
            if field_name in data:
                if data[field_name]:
                    try:
                        date_value = datetime.strptime(data[field_name], '%Y-%m-%d').date()
                        setattr(proyecto, model_field, date_value)
                    except (ValueError, TypeError):
                        pass
                else:
                    setattr(proyecto, model_field, None)

        # Cómputo
        if 'Computo' in data:
            proyecto.computo = data['Computo']

        # Campos pendientes
        pendiente_fields = {
            'NTP': 'ntp',
            'SCAN': 'scan',
            'RESUELVE POR NOMBRE': 'resuelve_por_nombre',
            'Antivirus': 'antivirus',  # clave alineada con el frontend
            'CONFIG BACKUP': 'config_backup',
            'MONITOREO NAGIOS': 'monitoreo_nagios',
            'MONITOREO ELASTIC': 'monitoreo_elastic',
            'UCMDB': 'ucmdb',
            'CONECTIVIDAD AWX 172.18.90.250 (SOLO UNIX)': 'conectividad_awx',
            'CAMBIO PASO OPERACIÓN (OLA)': 'cambio_paso_operacion_ola',
        }

        for excel_field, model_field in pendiente_fields.items():
            if excel_field in data:
                value = data[excel_field]
                if not value or (isinstance(value, str) and value.strip() == ''):
                    value = 'Pendiente'
                setattr(proyecto, model_field, value)

        # Otros campos
        if 'Base de Datos' in data:
            proyecto.base_de_datos = data['Base de Datos']
        if 'Balanceo' in data:
            proyecto.balanceo = data['Balanceo']
        if 'Backup' in data:
            proyecto.backup = data['Backup']
        if 'CHECK AV' in data:
            proyecto.check_av = data['CHECK AV']

        # Campos adicionales
        campos_adicionales = {
            'CONTACTO': 'contacto_id',  # Agregar campo CONTACTO
            'CANTIDAD MAQUINAS': 'cantidad_maquinas',
            'COD SERV_HOSTNAME': 'cod_serv_hostname',
            'PLATAFORMA': 'plataforma',
            'SO': 'so',
            'WINDOWS LICENCIA ACTIVADA': 'windows_licencia_activada',
            'DOMINIO': 'dominio',
            'PLATAFORMA BACKUP': 'plataforma_backup',
            'PROVEEDOR': 'proveedor',
            'COMUNIDAD SNMP': 'comunidad_snmp',
            'FGN 172.22.16.93': 'fgn_172_22_16_93',
            'RT': 'rt',
            'SERVICIO': 'servicio',
            'OBSERVACIONES': 'observaciones',
        }

        for excel_field, model_field in campos_adicionales.items():
            if excel_field in data:
                setattr(proyecto, model_field, data[excel_field] or None)

        proyecto.save()

        # Refrescar relación contacto para asegurar datos actualizados
        if proyecto.contacto_id:
            proyecto.refresh_from_db(fields=['contacto'])

        # 2. Requerimiento: Actualizar Fase si se envía en el JSON
        nueva_fase = data.get('Nueva Fase')
        fecha_fase_str = data.get('Fecha Fase')

        if nueva_fase and fecha_fase_str:
            try:
                fecha_fase = datetime.strptime(fecha_fase_str, '%Y-%m-%d').date()

                # Buscar si ya existe un registro para esta fase (tomar el más reciente si hay varios)
                existente = ProyectoFase.objects.filter(
                    proyecto=proyecto,
                    fase=nueva_fase
                ).order_by('-fecha', '-created_at').first()

                if existente:
                    # Si la fecha es distinta, actualizar la fecha de ese registro existente
                    if existente.fecha != fecha_fase:
                        try:
                            existente.fecha = fecha_fase
                            existente.save()
                        except Exception:
                            # Si hay conflicto de unicidad u otro error, crear una nueva fila como fallback
                            try:
                                ProyectoFase.objects.create(
                                    proyecto=proyecto,
                                    fase=nueva_fase,
                                    fecha=fecha_fase
                                )
                            except Exception as e:
                                print(f"Error actualizando/creando fase existente: {e}")
                    # si la fecha es la misma, no hacer nada (evitar duplicados)
                else:
                    # No existía la fase: crear nueva fila en el historial
                    try:
                        ProyectoFase.objects.create(
                            proyecto=proyecto,
                            fase=nueva_fase,
                            fecha=fecha_fase
                        )
                    except Exception as e:
                        print(f"Error creando nueva fase en historial: {e}")
            except Exception as e:
                print(f"Error actualizando fase: {e}")

        # Obtener la fase más reciente para la respuesta
        fase_actual = proyecto.fases.order_by('-fecha', '-created_at').first()
        if fase_actual:
            fase_str = f"{fase_actual.get_fase_display()} ({fase_actual.fecha.strftime('%Y-%m-%d')})"
        elif proyecto.estado == 'En Curso':
            fase_str = "Despliegue (No registrado)"
        else:
            fase_str = "Sin Fase"

        # Obtener información del contacto
        contacto_info = ''
        if proyecto.contacto:
            contacto_info = proyecto.contacto.nombre
            if proyecto.contacto.correo:
                contacto_info += f" ({proyecto.contacto.correo})"

        # Retornar el proyecto actualizado
        proyecto_dict = {
            'Id Project': proyecto.id_project,
            'RF': proyecto.rf or '',
            'Project': proyecto.project or '',
            'Project Leader': proyecto.project_leader or '',
            'Estado': proyecto.estado or '',
            'Fase': fase_str,
            'Start': proyecto.start.strftime('%Y-%m-%d') if proyecto.start else None,
            'Finish': proyecto.finish.strftime('%Y-%m-%d') if proyecto.finish else None,
            'Computo': proyecto.computo or '',
            'NTP': proyecto.ntp or '',
            'SCAN': proyecto.scan or '',
            'RESUELVE POR NOMBRE': proyecto.resuelve_por_nombre or '',
            'Antivirus': proyecto.antivirus or '',
            'CONFIG BACKUP': proyecto.config_backup or '',
            'MONITOREO NAGIOS': proyecto.monitoreo_nagios or '',
            'MONITOREO ELASTIC': proyecto.monitoreo_elastic or '',
            'UCMDB': proyecto.ucmdb or '',
            'CONECTIVIDAD AWX 172.18.90.250 (SOLO UNIX)': proyecto.conectividad_awx or '',
            'CAMBIO PASO OPERACIÓN (OLA)': proyecto.cambio_paso_operacion_ola or '',
            'Base de Datos': proyecto.base_de_datos or '',
            'Balanceo': proyecto.balanceo or '',
            'Backup': proyecto.backup or '',
            'CHECK AV': proyecto.check_av or '',
            'CANTIDAD MAQUINAS': proyecto.cantidad_maquinas or '',
            'COD SERV_HOSTNAME': proyecto.cod_serv_hostname or '',
            'PLATAFORMA': proyecto.plataforma or '',
            'SO': proyecto.so or '',
            'WINDOWS LICENCIA ACTIVADA': proyecto.windows_licencia_activada or '',
            'DOMINIO': proyecto.dominio or '',
            'PLATAFORMA BACKUP': proyecto.plataforma_backup or '',
            'PROVEEDOR': proyecto.proveedor or '',
            'COMUNIDAD SNMP': proyecto.comunidad_snmp or '',
            'FGN 172.22.16.93': proyecto.fgn_172_22_16_93 or '',
            'RT': proyecto.rt or '',
            'SERVICIO': proyecto.servicio or '',
            'OBSERVACIONES': proyecto.observaciones or '',
            'CONTACTO': contacto_info,
            'CONTACTO_ID': proyecto.contacto.id if proyecto.contacto else None,
        }

        return JsonResponse(proyecto_dict)

    except Proyecto.DoesNotExist:
        return JsonResponse({'error': 'Proyecto no encontrado'}, status=404)
    except Exception as e:
        return JsonResponse({'error': f'No se pudo actualizar el proyecto: {str(e)}'}, status=500)


@require_http_methods(["PUT"])
@login_required
def api_projects_update_status(request, project_id):
    """Actualizar el estado de un campo específico de un proyecto"""
    if request.method != 'PUT':
        return JsonResponse({'error': 'Método no permitido'}, status=405)

    try:
        proyecto = Proyecto.objects.get(id_project=project_id)
        data = json.loads(request.body)

        field_name = data.get('field_name')
        new_status = data.get('new_status')

        if not field_name or not new_status:
            return JsonResponse({'error': 'Faltan field_name o new_status'}, status=400)

        # Mapear nombres de campos de Excel a campos del modelo
        field_mapping = {
            'NTP': 'ntp',
            'SCAN': 'scan',
            'RESUELVE POR NOMBRE': 'resuelve_por_nombre',
            'Antivirus': 'antivirus',
            'CONFIG BACKUP': 'config_backup',
            'MONITOREO NAGIOS': 'monitoreo_nagios',
            'MONITOREO ELASTIC': 'monitoreo_elastic',
            'UCMDB': 'ucmdb',
            'CONECTIVIDAD AWX 172.18.90.250 (SOLO UNIX)': 'conectividad_awx',
            'CAMBIO PASO OPERACIÓN (OLA)': 'cambio_paso_operacion_ola',
            'WINDOWS LICENCIA ACTIVADA': 'windows_licencia_activada',
        }

        # Buscar el campo en el mapeo (case-insensitive)
        model_field = None
        for excel_field, model_field_name in field_mapping.items():
            if excel_field.lower() == field_name.lower():
                model_field = model_field_name
                break

        if not model_field or not hasattr(proyecto, model_field):
            return JsonResponse({'error': f'Campo "{field_name}" no encontrado en el proyecto'}, status=404)

        setattr(proyecto, model_field, new_status)
        proyecto.save()

        return JsonResponse({'success': True, 'message': 'Estado del proyecto actualizado correctamente'})

    except Proyecto.DoesNotExist:
        return JsonResponse({'error': 'Proyecto no encontrado'}, status=404)
    except Exception as e:
        return JsonResponse({'error': f'No se pudo actualizar el estado del proyecto: {str(e)}'}, status=500)


# ==================== API DE EVENTOS ====================

@require_http_methods(["GET", "POST"])
@login_required
def api_events(request):
    """Obtener todos los eventos o crear uno nuevo"""
    if request.method == 'GET':
        eventos = Evento.objects.all().order_by('fecha_inicio')

        eventos_list = []
        for e in eventos:
            evento_dict = {
                'id': e.id,
                'Titulo': e.titulo or '',
                'Descripcion': e.descripcion or '',
                'Fecha de Inicio': e.fecha_inicio.isoformat() if e.fecha_inicio else None,
                'Fecha de Fin': e.fecha_fin.isoformat() if e.fecha_fin else None,
                'Ubicacion': e.ubicacion or '',
                'Responsable': e.responsable or '',
            }
            eventos_list.append(evento_dict)

        return JsonResponse(eventos_list, safe=False)

    elif request.method == 'POST':
        """Agregar un nuevo evento"""
        try:
            data = json.loads(request.body)

            evento = Evento()
            evento.titulo = data.get('Titulo', '')
            evento.descripcion = data.get('Descripcion', '')
            evento.ubicacion = data.get('Ubicacion', '')
            evento.responsable = data.get('Responsable', '')

            # Fechas
            if 'Fecha de Inicio' in data and data['Fecha de Inicio']:
                try:
                    evento.fecha_inicio = datetime.fromisoformat(data['Fecha de Inicio'].replace('Z', '+00:00'))
                except (ValueError, TypeError):
                    pass

            if 'Fecha de Fin' in data and data['Fecha de Fin']:
                try:
                    evento.fecha_fin = datetime.fromisoformat(data['Fecha de Fin'].replace('Z', '+00:00'))
                except (ValueError, TypeError):
                    pass

            evento.save()

            evento_dict = {
                'id': evento.id,
                'Titulo': evento.titulo or '',
                'Descripcion': evento.descripcion or '',
                'Fecha de Inicio': evento.fecha_inicio.isoformat() if evento.fecha_inicio else None,
                'Fecha de Fin': evento.fecha_fin.isoformat() if evento.fecha_fin else None,
                'Ubicacion': evento.ubicacion or '',
                'Responsable': evento.responsable or '',
            }

            return JsonResponse(evento_dict, status=201)

        except Exception as e:
            return JsonResponse({'error': f'No se pudo guardar el nuevo evento: {str(e)}'}, status=500)

    else:
        return JsonResponse({'error': 'Método no permitido'}, status=405)


@require_http_methods(["GET"])
@login_required
def api_events_next(request):
    """Obtener el próximo evento"""
    now = timezone.now()
    next_event = Evento.objects.filter(fecha_inicio__gt=now).order_by('fecha_inicio').first()

    if not next_event:
        return JsonResponse({'message': 'No hay eventos próximos'}, status=404)

    evento_dict = {
        'id': next_event.id,
        'Titulo': next_event.titulo or '',
        'Descripcion': next_event.descripcion or '',
        'Fecha de Inicio': next_event.fecha_inicio.isoformat() if next_event.fecha_inicio else None,
        'Fecha de Fin': next_event.fecha_fin.isoformat() if next_event.fecha_fin else None,
        'Ubicacion': next_event.ubicacion or '',
        'Responsable': next_event.responsable or '',
    }

    return JsonResponse(evento_dict)


@require_http_methods(["POST"])
@login_required
def api_events_add(request):
    """Agregar un nuevo evento"""
    try:
        data = json.loads(request.body)

        evento = Evento()
        evento.titulo = data.get('Titulo', '')
        evento.descripcion = data.get('Descripcion', '')
        evento.ubicacion = data.get('Ubicacion', '')
        evento.responsable = data.get('Responsable', '')

        # Fechas
        if 'Fecha de Inicio' in data and data['Fecha de Inicio']:
            try:
                evento.fecha_inicio = datetime.fromisoformat(data['Fecha de Inicio'].replace('Z', '+00:00'))
            except (ValueError, TypeError):
                pass

        if 'Fecha de Fin' in data and data['Fecha de Fin']:
            try:
                evento.fecha_fin = datetime.fromisoformat(data['Fecha de Fin'].replace('Z', '+00:00'))
            except (ValueError, TypeError):
                pass

        evento.save()

        evento_dict = {
            'id': evento.id,
            'Titulo': evento.titulo or '',
            'Descripcion': evento.descripcion or '',
            'Fecha de Inicio': evento.fecha_inicio.isoformat() if evento.fecha_inicio else None,
            'Fecha de Fin': evento.fecha_fin.isoformat() if evento.fecha_fin else None,
            'Ubicacion': evento.ubicacion or '',
            'Responsable': evento.responsable or '',
        }

        return JsonResponse(evento_dict, status=201)

    except Exception as e:
        return JsonResponse({'error': f'No se pudo guardar el nuevo evento: {str(e)}'}, status=500)


@require_http_methods(["PUT", "DELETE"])
@login_required
def api_events_update(request, event_id):
    """Actualizar o eliminar un evento existente"""
    if request.method == 'DELETE':
        """Eliminar un evento"""
        try:
            evento = Evento.objects.get(id=event_id)
            evento.delete()
            return JsonResponse({'success': True})

        except Evento.DoesNotExist:
            return JsonResponse({'error': 'Evento no encontrado'}, status=404)
        except Exception as e:
            return JsonResponse({'error': f'No se pudo eliminar el evento: {str(e)}'}, status=500)

    elif request.method == 'PUT':
        """Actualizar un evento existente"""
        try:
            evento = Evento.objects.get(id=event_id)
            data = json.loads(request.body)

            if 'Titulo' in data:
                evento.titulo = data['Titulo']
            if 'Descripcion' in data:
                evento.descripcion = data['Descripcion']
            if 'Ubicacion' in data:
                evento.ubicacion = data['Ubicacion']
            if 'Responsable' in data:
                evento.responsable = data['Responsable']

            # Fechas
            if 'Fecha de Inicio' in data and data['Fecha de Inicio']:
                try:
                    evento.fecha_inicio = datetime.fromisoformat(data['Fecha de Inicio'].replace('Z', '+00:00'))
                except (ValueError, TypeError):
                    pass

            if 'Fecha de Fin' in data and data['Fecha de Fin']:
                try:
                    evento.fecha_fin = datetime.fromisoformat(data['Fecha de Fin'].replace('Z', '+00:00'))
                except (ValueError, TypeError):
                    pass

            evento.save()

            evento_dict = {
                'id': evento.id,
                'Titulo': evento.titulo or '',
                'Descripcion': evento.descripcion or '',
                'Fecha de Inicio': evento.fecha_inicio.isoformat() if evento.fecha_inicio else None,
                'Fecha de Fin': evento.fecha_fin.isoformat() if evento.fecha_fin else None,
                'Ubicacion': evento.ubicacion or '',
                'Responsable': evento.responsable or '',
            }

            return JsonResponse(evento_dict)

        except Evento.DoesNotExist:
            return JsonResponse({'error': 'Evento no encontrado'}, status=404)
        except Exception as e:
            return JsonResponse({'error': f'No se pudo actualizar el evento: {str(e)}'}, status=500)

    else:
        return JsonResponse({'error': 'Método no permitido'}, status=405)


# ==================== API DE CONTACTOS ====================

@login_required
@require_http_methods(["GET"])
def api_contacts_simple(request):
    """Obtener contactos en formato simple para selector desplegable."""
    contactos = Contacto.objects.all().order_by('nombre')
    contactos_list = []
    for c in contactos:
        contactos_list.append({
            'id': c.id,
            'nombre': c.nombre,
            'correo': c.correo,
            'telefono': c.telefono,
        })
    return JsonResponse(contactos_list, safe=False)


@login_required
@require_http_methods(["GET"])
def api_inventario_all(request):
    """Obtener todo el inventario de equipos de todos los proyectos."""
    try:
        inventario = ControlProyectosInventario.objects.select_related('proyecto').all().order_by('proyecto__id_project', 'hostname', 'tipo_equipo')
        inventario_list = []
        for item in inventario:
            inventario_list.append({
                'id': item.id,
                'proyecto_id': item.proyecto.id_project if item.proyecto else None,
                'proyecto_nombre': item.proyecto.project if item.proyecto else '',
                'ubicacion': item.ubicacion,
                'ot': item.ot,
                'codigo': item.codigo,
                'hostname': item.hostname,
                'tipo_equipo': item.tipo_equipo,
                'cpu': item.cpu,
                'ram': item.ram,
                'disco_so': item.disco_so,
                'disco_pag': item.disco_pag,
                'disco_data': item.disco_data,
                'ip_gestion': item.ip_gestion,
                'ip_servicios': item.ip_servicios,
                'ip_produccion': item.ip_produccion,
                'ip_adicional_1': item.ip_adicional_1,
                'ip_adicional_2': item.ip_adicional_2,
                'sistema_operativo': item.sistema_operativo,
                'referencia': item.referencia,
                'created_at': item.created_at.isoformat() if item.created_at else None,
                'updated_at': item.updated_at.isoformat() if item.updated_at else None,
            })
        return JsonResponse(inventario_list, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_http_methods(["GET", "POST"])
def api_inventario(request):
    """Obtener inventario de equipos de un proyecto o crear uno nuevo."""
    if request.method == 'GET':
        proyecto_id = request.GET.get('proyecto_id')
        if not proyecto_id:
            return JsonResponse({'error': 'Se requiere proyecto_id'}, status=400)

        try:
            proyecto = Proyecto.objects.get(id_project=proyecto_id)
            inventario = ControlProyectosInventario.objects.filter(proyecto=proyecto).order_by('hostname', 'tipo_equipo')
            inventario_list = []
            for item in inventario:
                inventario_list.append({
                    'id': item.id,
                    'ubicacion': item.ubicacion,
                    'ot': item.ot,
                    'codigo': item.codigo,
                    'hostname': item.hostname,
                    'cpu': item.cpu,
                    'ram': item.ram,
                    'disco_so': item.disco_so,
                    'disco_pag': item.disco_pag,
                    'disco_data': item.disco_data,
                    'ip_gestion': item.ip_gestion,
                    'ip_servicios': item.ip_servicios,
                    'ip_produccion': item.ip_produccion,
                    'ip_adicional_1': item.ip_adicional_1,
                    'ip_adicional_2': item.ip_adicional_2,
                    'sistema_operativo': item.sistema_operativo,
                    'tipo_equipo': item.tipo_equipo,
                    'referencia': item.referencia,
                    'proyecto_id': item.proyecto.id_project,
                })
            return JsonResponse(inventario_list, safe=False)
        except Proyecto.DoesNotExist:
            return JsonResponse({'error': 'Proyecto no encontrado'}, status=404)

    elif request.method == 'POST':
        try:
            data = json.loads(request.body)
            proyecto_id = data.get('proyecto_id')

            if not proyecto_id:
                return JsonResponse({'error': 'Se requiere proyecto_id'}, status=400)

            try:
                proyecto = Proyecto.objects.get(id_project=proyecto_id)
            except Proyecto.DoesNotExist:
                return JsonResponse({'error': 'Proyecto no encontrado'}, status=404)

            inventario_item = ControlProyectosInventario.objects.create(
                ubicacion=data.get('ubicacion'),
                ot=data.get('ot'),
                codigo=data.get('codigo'),
                hostname=data.get('hostname'),
                cpu=data.get('cpu'),
                ram=data.get('ram'),
                disco_so=data.get('disco_so'),
                disco_pag=data.get('disco_pag'),
                disco_data=data.get('disco_data'),
                ip_gestion=data.get('ip_gestion'),
                ip_servicios=data.get('ip_servicios'),
                ip_produccion=data.get('ip_produccion'),
                ip_adicional_1=data.get('ip_adicional_1'),
                ip_adicional_2=data.get('ip_adicional_2'),
                sistema_operativo=data.get('sistema_operativo'),
                tipo_equipo=data.get('tipo_equipo'),
                referencia=data.get('referencia'),
                proyecto=proyecto,
            )

            return JsonResponse({
                'id': inventario_item.id,
                'ubicacion': inventario_item.ubicacion,
                'ot': inventario_item.ot,
                'codigo': inventario_item.codigo,
                'hostname': inventario_item.hostname,
                'cpu': inventario_item.cpu,
                'ram': inventario_item.ram,
                'disco_so': inventario_item.disco_so,
                'disco_pag': inventario_item.disco_pag,
                'disco_data': inventario_item.disco_data,
                'ip_gestion': inventario_item.ip_gestion,
                'ip_servicios': inventario_item.ip_servicios,
                'ip_produccion': inventario_item.ip_produccion,
                'ip_adicional_1': inventario_item.ip_adicional_1,
                'ip_adicional_2': inventario_item.ip_adicional_2,
                'sistema_operativo': inventario_item.sistema_operativo,
                'tipo_equipo': inventario_item.tipo_equipo,
                'referencia': inventario_item.referencia,
                'proyecto_id': inventario_item.proyecto.id_project,
            })
        except json.JSONDecodeError:
            return JsonResponse({'error': 'JSON inválido'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_http_methods(["GET", "PUT", "DELETE"])
def api_inventario_detail(request, pk):
    """Obtener, actualizar o eliminar un item del inventario."""
    try:
        item = ControlProyectosInventario.objects.get(pk=pk)
    except ControlProyectosInventario.DoesNotExist:
        return JsonResponse({'error': 'Item no encontrado'}, status=404)

    if request.method == 'GET':
        return JsonResponse({
            'id': item.id,
            'ubicacion': item.ubicacion,
            'ot': item.ot,
            'codigo': item.codigo,
            'hostname': item.hostname,
            'cpu': item.cpu,
            'ram': item.ram,
            'disco_so': item.disco_so,
            'disco_pag': item.disco_pag,
            'disco_data': item.disco_data,
            'ip_gestion': item.ip_gestion,
            'ip_servicios': item.ip_servicios,
            'ip_produccion': item.ip_produccion,
            'ip_adicional_1': item.ip_adicional_1,
            'ip_adicional_2': item.ip_adicional_2,
            'sistema_operativo': item.sistema_operativo,
            'tipo_equipo': item.tipo_equipo,
            'referencia': item.referencia,
            'proyecto_id': item.proyecto.id_project,
        })

    elif request.method == 'PUT':
        try:
            data = json.loads(request.body)

            item.ubicacion = data.get('ubicacion', item.ubicacion)
            item.ot = data.get('ot', item.ot)
            item.codigo = data.get('codigo', item.codigo)
            item.hostname = data.get('hostname', item.hostname)
            item.cpu = data.get('cpu', item.cpu)
            item.ram = data.get('ram', item.ram)
            item.disco_so = data.get('disco_so', item.disco_so)
            item.disco_pag = data.get('disco_pag', item.disco_pag)
            item.disco_data = data.get('disco_data', item.disco_data)
            item.ip_gestion = data.get('ip_gestion', item.ip_gestion)
            item.ip_servicios = data.get('ip_servicios', item.ip_servicios)
            item.ip_produccion = data.get('ip_produccion', item.ip_produccion)
            item.ip_adicional_1 = data.get('ip_adicional_1', item.ip_adicional_1)
            item.ip_adicional_2 = data.get('ip_adicional_2', item.ip_adicional_2)
            item.sistema_operativo = data.get('sistema_operativo', item.sistema_operativo)
            item.tipo_equipo = data.get('tipo_equipo', item.tipo_equipo)
            item.referencia = data.get('referencia', item.referencia)

            item.save()

            return JsonResponse({
                'id': item.id,
                'ubicacion': item.ubicacion,
                'ot': item.ot,
                'codigo': item.codigo,
                'hostname': item.hostname,
                'cpu': item.cpu,
                'ram': item.ram,
                'disco_so': item.disco_so,
                'disco_pag': item.disco_pag,
                'disco_data': item.disco_data,
                'ip_gestion': item.ip_gestion,
                'ip_servicios': item.ip_servicios,
                'ip_produccion': item.ip_produccion,
                'ip_adicional_1': item.ip_adicional_1,
                'ip_adicional_2': item.ip_adicional_2,
                'sistema_operativo': item.sistema_operativo,
                'tipo_equipo': item.tipo_equipo,
                'referencia': item.referencia,
                'proyecto_id': item.proyecto.id_project,
            })
        except json.JSONDecodeError:
            return JsonResponse({'error': 'JSON inválido'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    elif request.method == 'DELETE':
        item.delete()
        return JsonResponse({'success': True})


@login_required
@require_http_methods(["GET", "POST"])
def api_contacts(request):
    """Obtener todos los contactos o crear uno nuevo."""
    if request.method == 'GET':
        # Obtener contactos con los proyectos donde son contacto principal
        contactos = Contacto.objects.prefetch_related('proyectos_principales').all().order_by('nombre')
        contactos_list = []
        for c in contactos:
            # Obtener proyectos donde este contacto es el contacto principal
            proyectos = c.proyectos_principales.all()
            proyectos_data = [
                {'id': p.id_project, 'nombre': p.project}
                for p in proyectos
            ]
            contactos_list.append({
                'id': c.id,
                'nombre': c.nombre,
                'telefono': c.telefono,
                'correo': c.correo,
                'cargo': c.cargo,
                'area': c.area,
                'notas': c.notas,
                'proyectos': proyectos_data,
                'proyecto_id': proyectos[0].id_project if proyectos else None,
                'proyecto_nombre': proyectos[0].project if proyectos else '',
            })
        return JsonResponse(contactos_list, safe=False)

    elif request.method == 'POST':
        try:
            data = json.loads(request.body)

            contacto = Contacto.objects.create(
                nombre=data.get('nombre'),
                telefono=data.get('telefono'),
                correo=data.get('correo'),
                cargo=data.get('cargo'),
                area=data.get('area'),
                notas=data.get('notas'),
            )

            # Si se proporciona proyecto_id, asignar el contacto al proyecto
            proyecto_id = data.get('proyecto_id')
            if proyecto_id:
                try:
                    proyecto = Proyecto.objects.get(id_project=proyecto_id)
                    proyecto.contacto = contacto
                    proyecto.save()
                except Proyecto.DoesNotExist:
                    pass

            return JsonResponse({
                'id': contacto.id,
                'nombre': contacto.nombre,
                'telefono': contacto.telefono,
                'correo': contacto.correo,
                'cargo': contacto.cargo,
                'area': contacto.area,
                'notas': contacto.notas,
                'proyectos': [],
                'proyecto_id': proyecto_id,
                'proyecto_nombre': '',
            }, status=201)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

@login_required
@require_http_methods(["PUT", "DELETE"])
def api_contact_detail(request, contact_id):
    """Actualizar o eliminar un contacto."""
    try:
        contacto = Contacto.objects.prefetch_related('proyectos_principales').get(pk=contact_id)
    except Contacto.DoesNotExist:
        return JsonResponse({'error': 'Contacto no encontrado'}, status=404)

    if request.method == 'PUT':
        try:
            data = json.loads(request.body)
            for field in ['nombre', 'telefono', 'correo', 'cargo', 'area', 'notas']:
                if field in data:
                    setattr(contacto, field, data[field])

            # Si se proporciona proyecto_id, asignar este contacto al proyecto
            if 'proyecto_id' in data:
                proyecto_id = data.get('proyecto_id')
                if proyecto_id:
                    try:
                        proyecto = Proyecto.objects.get(id_project=proyecto_id)
                        proyecto.contacto = contacto
                        proyecto.save()
                    except Proyecto.DoesNotExist:
                        pass

            contacto.save()

            # Obtener proyectos donde este contacto es el contacto principal
            proyectos = contacto.proyectos_principales.all()
            proyecto_id = proyectos[0].id_project if proyectos else None
            proyecto_nombre = proyectos[0].project if proyectos else ''

            return JsonResponse({
                'id': contacto.id,
                'nombre': contacto.nombre,
                'telefono': contacto.telefono,
                'correo': contacto.correo,
                'cargo': contacto.cargo,
                'area': contacto.area,
                'notas': contacto.notas,
                'proyecto_id': proyecto_id,
                'proyecto_nombre': proyecto_nombre,
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    elif request.method == 'DELETE':
        contacto.delete()
        return JsonResponse({'success': True}, status=204)


# ==================== GENERACIÓN DE INFORMES CON IA ====================

@login_required
def generar_informe_ia(request):
    """Generar informe con IA usando Google Gemini"""
    is_xhr = request.GET.get('xhr') == '1' or request.headers.get('X-Requested-With') == 'XMLHttpRequest'

    api_key = settings.GEMINI_API_KEY
    if not api_key:
        if is_xhr:
            return JsonResponse({'success': False, 'message': 'La clave de API de Gemini no está configurada.'})
        messages.error(request, "La clave de API de Gemini no está configurada.")
        return redirect('index')

    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash-lite-001')
    except Exception as e:
        if is_xhr:
            return JsonResponse({'success': False, 'message': f'Error al configurar la API de Gemini: {str(e)}'})
        messages.error(request, f"Error al configurar la API de Gemini: {e}")
        return redirect('index')

    # Obtener proyectos en curso
    proyectos_en_curso = Proyecto.objects.filter(estado='En Curso')

    if not proyectos_en_curso.exists():
        if is_xhr:
            return JsonResponse({'success': False, 'message': "No hay proyectos 'En Curso' para generar un informe."})
        messages.info(request, "No hay proyectos 'En Curso' para generar un informe.")
        return redirect('index')

    # Procesar cada proyecto con IA
    informe_parts = []
    proyecto_count = len(proyectos_en_curso)
    for idx, proyecto in enumerate(proyectos_en_curso):
        # Delay entre proyectos para evitar rate limit (excepto el primero)
        if idx > 0:
            time.sleep(5)
        # Obtener inventario del proyecto
        inventario_items = proyecto.inventario_equipos.all()

        # Construir tabla de inventario
        if inventario_items:
            tabla_inventario = []
            tabla_inventario.append("| Hostname | Tipo | CPU | RAM | IP Gestión | IP Servicios | Sistema Operativo |")
            tabla_inventario.append("|----------|------|-----|-----|------------|--------------|-------------------|")
            for item in inventario_items:
                hostname = item.hostname or 'N/A'
                tipo = item.tipo_equipo or 'N/A'
                cpu = item.cpu or 'N/A'
                ram = item.ram or 'N/A'
                ip_gestion = item.ip_gestion or 'N/A'
                ip_servicios = item.ip_servicios or 'N/A'
                so = item.sistema_operativo or 'N/A'
                tabla_inventario.append(f"| {hostname} | {tipo} | {cpu} | {ram} | {ip_gestion} | {ip_servicios} | {so} |")
            inventario_texto = "\n".join(tabla_inventario)
        else:
            inventario_texto = "No hay equipos registrados en el inventario."

        prompt_data = {
            "ID": proyecto.id_project or 'N/A',
            "RF": proyecto.rf or 'N/A',
            "Nombre del Proyecto": proyecto.project or 'N/A',
            "Estado": proyecto.estado or 'N/A',
            "Líder de Proyecto": proyecto.project_leader or 'N/A',
            "Inicio Real": proyecto.start.strftime('%Y-%m-%d') if proyecto.start else 'N/A',
            "Fin Real": proyecto.finish.strftime('%Y-%m-%d') if proyecto.finish else 'N/A',
            "Progreso": "N/A",
        }

        prompt = (
            "Eres un asistente experto en gestión de proyectos. "
            "Basado en los siguientes datos de un proyecto y su inventario de equipos, genera un 'Análisis de Estado' breve, profesional y accionable (2-3 frases). "
            "El análisis debe interpretar los datos clave (progreso, fechas, infraestructura) y describir la situación actual del proyecto, sugiriendo un siguiente paso o punto de atención. "
            "No incluyas un título, solo el párrafo del análisis.\n\n"
            "Datos del Proyecto:\n"
            f"- ID del Proyecto: {prompt_data['ID']}\n"
            f"- RF: {prompt_data['RF']}\n"
            f"- Nombre: {prompt_data['Nombre del Proyecto']}\n"
            f"- Progreso: {prompt_data['Progreso']}\n"
            f"- Estado: {prompt_data['Estado']}\n"
            f"- Líder: {prompt_data['Líder de Proyecto']}\n"
            f"- Fechas Reales (Inicio/Fin): {prompt_data['Inicio Real']} / {prompt_data['Fin Real']}\n\n"
            f"Inventario de Equipos ({len(inventario_items)} equipos):\n"
            f"{inventario_texto}\n\n"
            "Análisis de Estado:"
        )

        # Generar análisis con reintentos
        max_attempts = 3
        base_backoff = 1.0
        attempt = 0
        success = False
        last_error = None
        analisis = "No se pudo generar el análisis debido a un error."

        while attempt < max_attempts and not success:
            try:
                time.sleep(1)  # Rate limiting básico
                response = model.generate_content(prompt)
                analisis = response.text.strip().replace('*', '')
                success = True
            except Exception as e:
                last_error = e
                msg = str(e)
                if 'Resource exhausted' in msg or '429' in msg or 'rate limit' in msg.lower() or 'quota' in msg.lower():
                    attempt += 1
                    sleep_time = base_backoff * (2 ** (attempt - 1))
                    time.sleep(sleep_time)
                    continue
                else:
                    break

        if not success and last_error:
            msg = str(last_error)
            if 'Resource exhausted' in msg or '429' in msg or 'rate limit' in msg.lower() or 'quota' in msg.lower():
                analisis = "No se pudo generar el análisis: límite de recursos alcanzado. Intenta más tarde."

        title = proyecto.project or f"Proyecto {proyecto.id_project}"
        informe_parts.append(f"<section class='mb-3'><h3>{title}</h3><p>{analisis}</p></section>")

    informe_html = '\n'.join(informe_parts) if informe_parts else '<p>No hay contenido en el informe.</p>'

    if is_xhr:
        try:
            rendered = render(request, 'informe.html', {'informe': informe_html})
            return JsonResponse({'success': True, 'html': rendered.content.decode('utf-8')})
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error al renderizar el informe: {str(e)}'}, status=500)

    return render(request, 'informe.html', {'informe': informe_html})


@login_required
def generar_informe_tradicional(request):
    """Generar informe ejecutivo tradicional en HTML para impresión/PDF"""
    # Obtener proyectos en curso con contacto e inventario relacionado
    proyectos = Proyecto.objects.filter(estado='En Curso').select_related('contacto').prefetch_related('inventario_equipos').order_by('id_project')

    # Calcular estadísticas
    total_proyectos = proyectos.count()
    total_equipos = sum(p.inventario_equipos.count() for p in proyectos)
    proyectos_con_inventario = sum(1 for p in proyectos if p.inventario_equipos.exists())
    proyectos_sin_inventario = total_proyectos - proyectos_con_inventario

    context = {
        'proyectos': proyectos,
        'total_proyectos': total_proyectos,
        'total_equipos': total_equipos,
        'proyectos_con_inventario': proyectos_con_inventario,
        'proyectos_sin_inventario': proyectos_sin_inventario,
        'fecha_generacion': timezone.now().strftime('%d/%m/%Y %H:%M'),
        'usuario': request.user.get_full_name() or request.user.username,
    }

    return render(request, 'informe_tradicional.html', context)


@login_required
def inventory_general(request):
    """Vista para mostrar el inventario general de todos los proyectos."""
    return render(request, 'inventory/general.html', {
        'title': 'Inventario General',
        'description': 'Vista completa de todo el inventario de equipos'
    })


@login_required
def inventory_projects_in_progress(request):
    """Vista para mostrar el inventario de proyectos en curso."""
    return render(request, 'inventory/projects_in_progress.html', {
        'title': 'Inventario de Proyectos en Curso',
        'description': 'Equipos de proyectos que están actualmente en desarrollo o ejecución'
    })


@login_required
def exportar_informe_word(request):
    """Exportar informe ejecutivo a Word (DOCX)"""
    # Obtener proyectos en curso con contacto e inventario relacionado
    proyectos = Proyecto.objects.filter(estado='En Curso').select_related('contacto').prefetch_related('inventario_equipos').order_by('id_project')

    # Crear documento Word
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Título principal
    titulo = doc.add_heading('Informe Ejecutivo - Proyectos en Curso', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fecha del informe
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_run = fecha.add_run(f"Generado el: {timezone.now().strftime('%d/%m/%Y %H:%M')}")
    fecha_run.font.size = Pt(10)
    fecha_run.font.color.rgb = RGBColor(128, 128, 128)

    # Espacio
    doc.add_paragraph()

    # Resumen ejecutivo
    total_proyectos = proyectos.count()
    if total_proyectos > 0:
        resumen = doc.add_heading('Resumen Ejecutivo', level=1)

        p = doc.add_paragraph()
        p.add_run(f'Total de proyectos en curso: ').bold = True
        p.add_run(str(total_proyectos))

        # Contar proyectos con inventario
        proyectos_con_inventario = sum(1 for p in proyectos if p.inventario_equipos.exists())
        p = doc.add_paragraph()
        p.add_run(f'Proyectos con inventario: ').bold = True
        p.add_run(f'{proyectos_con_inventario} de {total_proyectos}')

        # Contar proyectos con contacto
        proyectos_con_contacto = sum(1 for p in proyectos if p.contacto)
        p = doc.add_paragraph()
        p.add_run(f'Proyectos con contacto asignado: ').bold = True
        p.add_run(f'{proyectos_con_contacto} de {total_proyectos}')

        doc.add_paragraph()

    # Detalle de proyectos
    doc.add_heading('Detalle de Proyectos', level=1)

    for proyecto in proyectos:
        # Título del proyecto
        project_heading = doc.add_heading(
            f"Proyecto: {proyecto.project} (ID: {proyecto.id_project})",
            level=2
        )

        # Información básica
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'

        # Llenar tabla de información básica
        info_data = [
            ('Estado:', proyecto.estado or 'N/A'),
            ('Líder de Proyecto:', proyecto.project_leader or 'N/A'),
            ('RF:', proyecto.rf or 'N/A'),
            ('Servicio:', proyecto.servicio or 'N/A'),
        ]

        for idx, (label, value) in enumerate(info_data):
            row = info_table.rows[idx]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = str(value)

        doc.add_paragraph()

        # Contacto
        if proyecto.contacto:
            contacto_heading = doc.add_heading('Contacto', level=3)
            contacto_para = doc.add_paragraph()
            contacto_para.add_run(f"Nombre: ").bold = True
            contacto_para.add_run(f"{proyecto.contacto.nombre}\n")
            contacto_para.add_run(f"Correo: ").bold = True
            contacto_para.add_run(f"{proyecto.contacto.correo or 'N/A'}\n")
            contacto_para.add_run(f"Teléfono: ").bold = True
            contacto_para.add_run(f"{proyecto.contacto.telefono or 'N/A'}\n")
            if proyecto.contacto.cargo:
                contacto_para.add_run(f"Cargo: ").bold = True
                contacto_para.add_run(f"{proyecto.contacto.cargo}")
        else:
            doc.add_heading('Contacto', level=3)
            p = doc.add_paragraph()
            p.add_run('Sin contacto asignado').italic = True

        doc.add_paragraph()

        # Inventario
        inventario = proyecto.inventario_equipos.all()
        if inventario:
            doc.add_heading('Inventario de Equipos', level=3)

            # Tabla de inventario
            inv_table = doc.add_table(rows=1, cols=5)
            inv_table.style = 'Medium Grid 1 Accent 1'

            # Encabezados
            headers = ['Tipo de Equipo', 'Hostname', 'IP Gestión', 'Ubicación', 'Sistema Operativo']
            header_cells = inv_table.rows[0].cells
            for idx, header in enumerate(headers):
                header_cells[idx].text = header
                header_cells[idx].paragraphs[0].runs[0].font.bold = True

            # Datos del inventario
            for item in inventario:
                row_cells = inv_table.add_row().cells
                row_cells[0].text = item.tipo_equipo or 'N/A'
                row_cells[1].text = item.hostname or 'N/A'
                row_cells[2].text = item.ip_gestion or 'N/A'
                row_cells[3].text = item.ubicacion or 'N/A'
                row_cells[4].text = item.sistema_operativo or 'N/A'
        else:
            doc.add_heading('Inventario de Equipos', level=3)
            p = doc.add_paragraph()
            p.add_run('No hay equipos registrados').italic = True

        # Separador entre proyectos
        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        doc.add_paragraph()

    # Guardar en buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Preparar respuesta
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="informe_ejecutivo_{timezone.now().strftime("%Y%m%d_%H%M")}.docx"'

    return response


@login_required
def inventory_projects_finished(request):
    """Vista para mostrar el inventario de proyectos finalizados y cerrados."""
    return render(request, 'inventory/projects_finished.html', {
        'title': 'Inventario de Proyectos Finalizados',
        'description': 'Equipos de proyectos que han sido finalizados y cerrados'
    })
